"""Document processor for handling .doc files and converting them to text chunks."""

import os
import logging
from typing import List, Dict, Any
from pathlib import Path

logger = logging.getLogger(__name__)

class DocumentProcessor:
    """Process .doc/.docx files and extract text for RAG system."""
    
    def __init__(self, docs_directory: str = None):
        """Initialize document processor.
        
        Args:
            docs_directory: Path to directory containing .doc files
        """
        if docs_directory is None:
            # Default to docs folder in yQi directory
            # Get absolute path to ensure it works from any calling location
            models_dir = Path(__file__).parent.absolute()
            yqi_root = models_dir.parent
            self.docs_directory = yqi_root / "docs"
        else:
            self.docs_directory = Path(docs_directory)
        
        self.supported_extensions = ['.doc', '.docx', '.txt', '.pdf']
    
    def extract_text_from_doc(self, file_path: str) -> str:
        """Extract text from .doc/.docx file.
        
        Args:
            file_path: Path to document file
            
        Returns:
            Extracted text content
        """
        try:
            file_path = Path(file_path)
            
            if file_path.suffix.lower() == '.txt':
                with open(file_path, 'r', encoding='utf-8') as f:
                    return f.read()
            
            elif file_path.suffix.lower() == '.docx':
                try:
                    # Use python-docx for .docx files
                    import docx
                    doc = docx.Document(file_path)
                    text = []
                    for paragraph in doc.paragraphs:
                        text.append(paragraph.text)
                    return '\n'.join(text)
                except ImportError:
                    logger.warning("python-docx not installed. Install with: pip install python-docx")
                    return ""
                except Exception as e:
                    logger.error(f"Error reading .docx file {file_path} with python-docx: {e}")
                    return ""
            
            elif file_path.suffix.lower() == '.doc':
                # Old .doc files require different handling
                try:
                    # Try textract first for .doc files
                    import textract
                    text = textract.process(str(file_path)).decode('utf-8')
                    return text
                except ImportError:
                    logger.warning("textract not installed. Trying alternative method...")
                    # Try using python-docx anyway (sometimes works)
                    try:
                        import docx
                        doc = docx.Document(file_path)
                        text = []
                        for paragraph in doc.paragraphs:
                            text.append(paragraph.text)
                        return '\n'.join(text)
                    except Exception as e2:
                        logger.error(f"Cannot process .doc file {file_path}. Install textract or convert to .docx: {e2}")
                        return ""
                except Exception as e:
                    logger.error(f"Error reading .doc file {file_path} with textract: {e}")
                    # Try python-docx as fallback
                    try:
                        import docx
                        doc = docx.Document(file_path)
                        text = []
                        for paragraph in doc.paragraphs:
                            text.append(paragraph.text)
                        return '\n'.join(text)
                    except Exception as e2:
                        logger.error(f"Fallback failed for .doc file {file_path}: {e2}")
                        return ""
            
            elif file_path.suffix.lower() == '.pdf':
                try:
                    # Try PyPDF2 first
                    import PyPDF2
                    with open(file_path, 'rb') as file:
                        reader = PyPDF2.PdfReader(file)
                        text = []
                        for page in reader.pages:
                            text.append(page.extract_text())
                        return '\n'.join(text)
                except ImportError:
                    logger.warning("PyPDF2 not installed. Trying textract for PDF...")
                    try:
                        import textract
                        text = textract.process(str(file_path)).decode('utf-8')
                        return text
                    except ImportError:
                        logger.warning("textract not installed. Install with: pip install PyPDF2 or textract")
                        return ""
                    except Exception as e:
                        logger.error(f"Error reading PDF {file_path} with textract: {e}")
                        return ""
                except Exception as e:
                    logger.error(f"Error reading PDF {file_path} with PyPDF2: {e}")
                    # Fallback to textract
                    try:
                        import textract
                        text = textract.process(str(file_path)).decode('utf-8')
                        return text
                    except ImportError:
                        logger.warning("textract not installed. Install with: pip install textract")
                        return ""
                    except Exception as e:
                        logger.error(f"Error reading PDF {file_path} with textract: {e}")
                        return ""
            
            else:
                logger.warning(f"Unsupported file format: {file_path.suffix}")
                return ""
                
        except Exception as e:
            logger.error(f"Error processing document {file_path}: {e}")
            return ""
    
    def chunk_text(self, text: str, chunk_size: int = 500, overlap: int = 50) -> List[str]:
        """Split text into overlapping chunks.
        
        Args:
            text: Text to chunk
            chunk_size: Maximum characters per chunk
            overlap: Number of characters to overlap between chunks
            
        Returns:
            List of text chunks
        """
        if not text.strip():
            return []
        
        chunks = []
        start = 0
        
        while start < len(text):
            end = start + chunk_size
            
            # Try to break at sentence boundaries
            if end < len(text):
                # Look for sentence endings within the last 100 characters
                sentence_end = text.rfind('.', start, end)
                if sentence_end > start + chunk_size - 100:
                    end = sentence_end + 1
                else:
                    # Look for paragraph breaks
                    para_break = text.rfind('\n\n', start, end)
                    if para_break > start + chunk_size - 200:
                        end = para_break + 2
            
            chunk = text[start:end].strip()
            if chunk:
                chunks.append(chunk)
            
            start = end - overlap
            
        return chunks
    
    def process_documents(self, chunk_size: int = 500, overlap: int = 50) -> List[Dict[str, Any]]:
        """Process all documents in the docs directory.
        
        Args:
            chunk_size: Maximum characters per chunk
            overlap: Number of characters to overlap between chunks
            
        Returns:
            List of document chunks with metadata
        """
        chunks = []
        
        # Debug: Print the path being checked
        logger.info(f"Looking for documents in: {self.docs_directory}")
        logger.info(f"Directory exists: {self.docs_directory.exists()}")
        
        if not self.docs_directory.exists():
            logger.warning(f"Documents directory not found: {self.docs_directory}")
            return chunks
        
        # Debug: List all files in directory
        try:
            all_files = list(self.docs_directory.iterdir())
            logger.info(f"All files in directory: {[f.name for f in all_files]}")
            
            supported_files = [f for f in all_files if f.is_file() and f.suffix.lower() in self.supported_extensions]
            logger.info(f"Supported files found: {[f.name for f in supported_files]}")
        except Exception as e:
            logger.error(f"Error listing directory contents: {e}")
            return chunks
        
        for file_path in self.docs_directory.iterdir():
            if file_path.is_file() and file_path.suffix.lower() in self.supported_extensions:
                logger.info(f"Processing document: {file_path.name}")
                
                try:
                    text = self.extract_text_from_doc(str(file_path))
                    logger.info(f"Extracted {len(text) if text else 0} characters from {file_path.name}")
                    
                    if text:
                        text_chunks = self.chunk_text(text, chunk_size, overlap)
                        logger.info(f"Created {len(text_chunks)} chunks from {file_path.name}")
                        
                        for i, chunk in enumerate(text_chunks):
                            chunks.append({
                                'text': chunk,
                                'source': file_path.name,
                                'chunk_id': f"{file_path.stem}_{i}",
                                'metadata': {
                                    'file_path': str(file_path),
                                    'chunk_index': i,
                                    'total_chunks': len(text_chunks)
                                }
                            })
                    else:
                        logger.warning(f"No text extracted from {file_path.name}")
                except Exception as e:
                    logger.error(f"Error processing {file_path.name}: {e}")
                    continue
        
        logger.info(f"Processed {len(chunks)} chunks from {len([f for f in self.docs_directory.iterdir() if f.suffix.lower() in self.supported_extensions])} documents")
        return chunks
    
    def get_available_documents(self) -> List[str]:
        """Get list of available document files.
        
        Returns:
            List of document filenames
        """
        if not self.docs_directory.exists():
            return []
        
        return [
            f.name for f in self.docs_directory.iterdir()
            if f.is_file() and f.suffix.lower() in self.supported_extensions
        ]
