"""Document processor for handling .doc files and converting them to text chunks."""

import os
import logging
from typing import List, Dict, Any
from pathlib import Path

logger = logging.getLogger(__name__)

class DocumentProcessor:
    """Process .doc/.docx files and extract text for RAG system."""
    
    def __init__(self, docs_directory: str = None):
        """Initialize document processor.
        
        Args:
            docs_directory: Path to directory containing .doc files
        """
        if docs_directory is None:
            # Default to docs folder in yQi directory
            # Get absolute path to ensure it works from any calling location
            models_dir = Path(__file__).parent.absolute()
            yqi_root = models_dir.parent
            self.docs_directory = yqi_root / "docs"
        else:
            self.docs_directory = Path(docs_directory)
        
        self.supported_extensions = ['.doc', '.docx', '.txt', '.pdf']
    
    def extract_text_from_doc(self, file_path: str) -> str:
        """Extract text from .doc/.docx file.
        
        Args:
            file_path: Path to document file
            
        Returns:
            Extracted text content
        """
        try:
            file_path = Path(file_path)
            
            if file_path.suffix.lower() == '.txt':
                with open(file_path, 'r', encoding='utf-8') as f:
                    return f.read()
            
            elif file_path.suffix.lower() == '.docx':
                try:
                    # Use python-docx for .docx files
                    import docx
                    doc = docx.Document(file_path)
                    text = []
                    for paragraph in doc.paragraphs:
                        text.append(paragraph.text)
                    return '\n'.join(text)
                except ImportError:
                    logger.warning("python-docx not installed. Install with: pip install python-docx")
                    return ""
                except Exception as e:
                    logger.error(f"Error reading .docx file {file_path} with python-docx: {e}")
                    return ""
            
            elif file_path.suffix.lower() == '.doc':
                # Old .doc files require different handling
                try:
                    # Try textract first for .doc files
                    import textract
                    text = textract.process(str(file_path)).decode('utf-8')
                    return text
                except ImportError:
                    logger.warning("textract not installed. Trying alternative method...")
                    # Try using python-docx anyway (sometimes works)
                    try:
                        import docx
                        doc = docx.Document(file_path)
                        text = []
                        for paragraph in doc.paragraphs:
                            text.append(paragraph.text)
                        return '\n'.join(text)
                    except Exception as e2:
                        logger.error(f"Cannot process .doc file {file_path}. Install textract or convert to .docx: {e2}")
                        return ""
                except Exception as e:
                    logger.error(f"Error reading .doc file {file_path} with textract: {e}")
                    # Try python-docx as fallback
                    try:
                        import docx
                        doc = docx.Document(file_path)
                        text = []
                        for paragraph in doc.paragraphs:
                            text.append(paragraph.text)
                        return '\n'.join(text)
                    except Exception as e2:
                        logger.error(f"Fallback failed for .doc file {file_path}: {e2}")
                        return ""
            
            elif file_path.suffix.lower() == '.pdf':
                try:
                    # Try PyPDF2 first
                    import PyPDF2
                    with open(file_path, 'rb') as file:
                        reader = PyPDF2.PdfReader(file)
                        text = []
                        for page in reader.pages:
                            text.append(page.extract_text())
                        return '\n'.join(text)
                except ImportError:
                    logger.warning("PyPDF2 not installed. Trying textract for PDF...")
                    try:
                        import textract
                        text = textract.process(str(file_path)).decode('utf-8')
                        return text
                    except ImportError:
                        logger.warning("textract not installed. Install with: pip install PyPDF2 or textract")
                        return ""
                    except Exception as e:
                        logger.error(f"Error reading PDF {file_path} with textract: {e}")
                        return ""
                except Exception as e:
                    logger.error(f"Error reading PDF {file_path} with PyPDF2: {e}")
                    # Fallback to textract
                    try:
                        import textract
                        text = textract.process(str(file_path)).decode('utf-8')
                        return text
                    except ImportError:
                        logger.warning("textract not installed. Install with: pip install textract")
                        return ""
                    except Exception as e:
                        logger.error(f"Error reading PDF {file_path} with textract: {e}")
                        return ""
            
            else:
                logger.warning(f"Unsupported file format: {file_path.suffix}")
                return ""
                
        except Exception as e:
            logger.error(f"Error processing document {file_path}: {e}")
            return ""
    
    def chunk_text(self, text: str, chunk_size: int = 500, overlap: int = 50) -> List[str]:
        """Split text into overlapping chunks.
        
        Args:
            text: Text to chunk
            chunk_size: Maximum characters per chunk
            overlap: Number of characters to overlap between chunks
            
        Returns:
            List of text chunks
        """
        if not text.strip():
            return []
        
        chunks = []
        start = 0
        
        while start < len(text):
            end = start + chunk_size
            
            # Try to break at sentence boundaries
            if end < len(text):
                # Look for sentence endings within the last 100 characters
                sentence_end = text.rfind('.', start, end)
                if sentence_end > start + chunk_size - 100:
                    end = sentence_end + 1
                else:
                    # Look for paragraph breaks
                    para_break = text.rfind('\n\n', start, end)
                    if para_break > start + chunk_size - 200:
                        end = para_break + 2
            
            chunk = text[start:end].strip()
            if chunk:
                chunks.append(chunk)
            
            start = end - overlap
            
        return chunks
    
    def detect_tcm_markers(self, text: str) -> Dict[str, List[int]]:
        """Detect TCM markers (#CHAPTER, #SECTION, #FORMULA) and their positions.
        
        Args:
            text: Text to analyze
            
        Returns:
            Dictionary with marker types as keys and lists of positions as values
        """
        markers = {
            'chapters': [],
            'sections': [],
            'formulas': []
        }
        
        lines = text.split('\n')
        char_position = 0
        
        for i, line in enumerate(lines):
            line_stripped = line.strip()
            
            if line_stripped == '#CHAPTER':
                markers['chapters'].append({
                    'position': char_position,
                    'line_number': i + 1,
                    'title': lines[i + 1].strip() if i + 1 < len(lines) else 'Unknown Chapter'
                })
            elif line_stripped == '#SECTION':
                markers['sections'].append({
                    'position': char_position,
                    'line_number': i + 1,
                    'title': lines[i + 1].strip() if i + 1 < len(lines) else 'Unknown Section'
                })
            elif line_stripped == '#FORMULA':
                markers['formulas'].append({
                    'position': char_position,
                    'line_number': i + 1,
                    'title': lines[i + 1].strip() if i + 1 < len(lines) else 'Unknown Formula'
                })
            
            char_position += len(line) + 1  # +1 for newline character
        
        return markers
    
    def chunk_by_chapters(self, text: str, max_chunk_size: int = 2000) -> List[Dict[str, Any]]:
        """Split text into chunks based on #CHAPTER markers.
        
        Args:
            text: Text to chunk
            max_chunk_size: Maximum size for each chunk (chapters may be split if too large)
            
        Returns:
            List of chunk dictionaries with metadata
        """
        if not text.strip():
            return []
        
        markers = self.detect_tcm_markers(text)
        chapters = markers['chapters']
        
        if not chapters:
            # No chapter markers found, fall back to regular chunking
            return [{'text': chunk, 'type': 'regular', 'title': 'No Chapter'} 
                   for chunk in self.chunk_text(text, max_chunk_size, 100)]
        
        chunks = []
        text_length = len(text)
        
        for i, chapter in enumerate(chapters):
            start_pos = chapter['position']
            end_pos = chapters[i + 1]['position'] if i + 1 < len(chapters) else text_length
            
            chapter_text = text[start_pos:end_pos].strip()
            chapter_title = chapter['title']
            
            # If chapter is too large, split it further
            if len(chapter_text) > max_chunk_size:
                # Try to split by sections within the chapter
                chapter_sections = []
                for section in markers['sections']:
                    if start_pos <= section['position'] < end_pos:
                        chapter_sections.append(section)
                
                if chapter_sections:
                    # Split by sections
                    for j, section in enumerate(chapter_sections):
                        section_start = section['position']
                        section_end = (chapter_sections[j + 1]['position'] 
                                     if j + 1 < len(chapter_sections) 
                                     else end_pos)
                        
                        section_text = text[section_start:section_end].strip()
                        if section_text:
                            chunks.append({
                                'text': section_text,
                                'type': 'section',
                                'chapter_title': chapter_title,
                                'section_title': section['title'],
                                'title': f"{chapter_title} - {section['title']}"
                            })
                else:
                    # No sections, split by regular chunking
                    chapter_chunks = self.chunk_text(chapter_text, max_chunk_size, 200)
                    for k, chunk_text in enumerate(chapter_chunks):
                        chunks.append({
                            'text': chunk_text,
                            'type': 'chapter_part',
                            'chapter_title': chapter_title,
                            'title': f"{chapter_title} (Part {k + 1})"
                        })
            else:
                # Chapter fits in one chunk
                chunks.append({
                    'text': chapter_text,
                    'type': 'chapter',
                    'chapter_title': chapter_title,
                    'title': chapter_title
                })
        
        return chunks
    
    def chunk_by_sections(self, text: str, max_chunk_size: int = 1500) -> List[Dict[str, Any]]:
        """Split text into chunks based on #SECTION markers.
        
        Args:
            text: Text to chunk
            max_chunk_size: Maximum size for each chunk
            
        Returns:
            List of chunk dictionaries with metadata
        """
        if not text.strip():
            return []
        
        markers = self.detect_tcm_markers(text)
        sections = markers['sections']
        chapters = markers['chapters']
        
        if not sections:
            # No section markers found, fall back to chapter chunking
            return self.chunk_by_chapters(text, max_chunk_size)
        
        chunks = []
        text_length = len(text)
        
        # Create a mapping of sections to their parent chapters
        section_to_chapter = {}
        for section in sections:
            parent_chapter = None
            for chapter in chapters:
                if chapter['position'] <= section['position']:
                    parent_chapter = chapter['title']
                else:
                    break
            section_to_chapter[section['position']] = parent_chapter or 'Unknown Chapter'
        
        for i, section in enumerate(sections):
            start_pos = section['position']
            end_pos = sections[i + 1]['position'] if i + 1 < len(sections) else text_length
            
            section_text = text[start_pos:end_pos].strip()
            section_title = section['title']
            chapter_title = section_to_chapter[section['position']]
            
            if len(section_text) > max_chunk_size:
                # Section is too large, split it further
                section_chunks = self.chunk_text(section_text, max_chunk_size, 100)
                for j, chunk_text in enumerate(section_chunks):
                    chunks.append({
                        'text': chunk_text,
                        'type': 'section_part',
                        'chapter_title': chapter_title,
                        'section_title': section_title,
                        'title': f"{chapter_title} - {section_title} (Part {j + 1})"
                    })
            else:
                # Section fits in one chunk
                chunks.append({
                    'text': section_text,
                    'type': 'section',
                    'chapter_title': chapter_title,
                    'section_title': section_title,
                    'title': f"{chapter_title} - {section_title}"
                })
        
        return chunks
    
    def process_documents_with_chunking(self, chunking_method: str = 'regular', 
                                      chunk_size: int = 500, overlap: int = 50) -> List[Dict[str, Any]]:
        """Process all documents with customizable chunking method.
        
        Args:
            chunking_method: 'regular', 'chapter', or 'section'
            chunk_size: Maximum characters per chunk (for regular chunking)
            overlap: Number of characters to overlap between chunks (for regular chunking)
            
        Returns:
            List of document chunks with metadata
        """
        chunks = []
        
        # Debug: Print the path being checked
        logger.info(f"Looking for documents in: {self.docs_directory}")
        logger.info(f"Directory exists: {self.docs_directory.exists()}")
        logger.info(f"Using chunking method: {chunking_method}")
        
        if not self.docs_directory.exists():
            logger.warning(f"Documents directory not found: {self.docs_directory}")
            return chunks
        
        # Debug: List all files in directory
        try:
            all_files = list(self.docs_directory.iterdir())
            logger.info(f"All files in directory: {[f.name for f in all_files]}")
            
            supported_files = [f for f in all_files if f.is_file() and f.suffix.lower() in self.supported_extensions]
            logger.info(f"Supported files found: {[f.name for f in supported_files]}")
        except Exception as e:
            logger.error(f"Error listing directory contents: {e}")
            return chunks
        
        for file_path in self.docs_directory.iterdir():
            if file_path.is_file() and file_path.suffix.lower() in self.supported_extensions:
                logger.info(f"Processing document: {file_path.name}")
                
                try:
                    text = self.extract_text_from_doc(str(file_path))
                    logger.info(f"Extracted {len(text) if text else 0} characters from {file_path.name}")
                    
                    if text:
                        # Choose chunking method
                        if chunking_method == 'chapter':
                            text_chunks = self.chunk_by_chapters(text, chunk_size)
                        elif chunking_method == 'section':
                            text_chunks = self.chunk_by_sections(text, chunk_size)
                        else:  # regular chunking
                            text_chunks = [{'text': chunk, 'type': 'regular', 'title': f'Chunk {i+1}'} 
                                         for i, chunk in enumerate(self.chunk_text(text, chunk_size, overlap))]
                        
                        logger.info(f"Created {len(text_chunks)} chunks from {file_path.name} using {chunking_method} method")
                        
                        for i, chunk_data in enumerate(text_chunks):
                            chunk_text = chunk_data['text'] if isinstance(chunk_data, dict) else chunk_data
                            chunk_title = chunk_data.get('title', f'Chunk {i+1}') if isinstance(chunk_data, dict) else f'Chunk {i+1}'
                            chunk_type = chunk_data.get('type', 'regular') if isinstance(chunk_data, dict) else 'regular'
                            
                            chunk_metadata = {
                                'file_path': str(file_path),
                                'chunk_index': i,
                                'total_chunks': len(text_chunks),
                                'chunking_method': chunking_method,
                                'chunk_type': chunk_type,
                                'chunk_title': chunk_title
                            }
                            
                            # Add chapter/section specific metadata if available
                            if isinstance(chunk_data, dict):
                                if 'chapter_title' in chunk_data:
                                    chunk_metadata['chapter_title'] = chunk_data['chapter_title']
                                if 'section_title' in chunk_data:
                                    chunk_metadata['section_title'] = chunk_data['section_title']
                            
                            chunks.append({
                                'text': chunk_text,
                                'source': file_path.name,
                                'chunk_id': f"{file_path.stem}_{i}",
                                'metadata': chunk_metadata
                            })
                    else:
                        logger.warning(f"No text extracted from {file_path.name}")
                except Exception as e:
                    logger.error(f"Error processing {file_path.name}: {e}")
                    continue
        
        logger.info(f"Processed {len(chunks)} chunks from {len([f for f in self.docs_directory.iterdir() if f.suffix.lower() in self.supported_extensions])} documents")
        return chunks
    
    def process_documents(self, chunk_size: int = 500, overlap: int = 50) -> List[Dict[str, Any]]:
        """Process all documents in the docs directory (legacy method for backward compatibility).
        
        Args:
            chunk_size: Maximum characters per chunk
            overlap: Number of characters to overlap between chunks
            
        Returns:
            List of document chunks with metadata
        """
        return self.process_documents_with_chunking('regular', chunk_size, overlap)
    
    def get_available_documents(self) -> List[str]:
        """Get list of available document files.
        
        Returns:
            List of document filenames
        """
        if not self.docs_directory.exists():
            return []
        
        return [
            f.name for f in self.docs_directory.iterdir()
            if f.is_file() and f.suffix.lower() in self.supported_extensions
        ]
    
    def get_supported_files(self) -> List[Path]:
        """Get list of supported document file paths.
        
        Returns:
            List of Path objects for supported files
        """
        if not self.docs_directory.exists():
            return []
        
        return [
            f for f in self.docs_directory.iterdir()
            if f.is_file() and f.suffix.lower() in self.supported_extensions
        ]
