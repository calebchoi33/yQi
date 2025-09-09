"""Text extraction utilities for different document formats."""

import logging
from abc import ABC, abstractmethod
from pathlib import Path
from typing import Optional

logger = logging.getLogger(__name__)


class TextExtractor(ABC):
    """Abstract base class for text extractors."""
    
    @abstractmethod
    def can_extract(self, file_path: Path) -> bool:
        """Check if this extractor can handle the given file."""
        pass
    
    @abstractmethod
    def extract_text(self, file_path: Path) -> str:
        """Extract text from the given file."""
        pass


class TxtExtractor(TextExtractor):
    """Text extractor for .txt files."""
    
    def can_extract(self, file_path: Path) -> bool:
        return file_path.suffix.lower() == '.txt'
    
    def extract_text(self, file_path: Path) -> str:
        """Extract text from .txt file."""
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                return f.read()
        except UnicodeDecodeError:
            # Try with different encodings
            for encoding in ['gbk', 'gb2312', 'latin-1']:
                try:
                    with open(file_path, 'r', encoding=encoding) as f:
                        return f.read()
                except UnicodeDecodeError:
                    continue
            logger.error(f"Could not decode {file_path} with any encoding")
            return ""
        except Exception as e:
            logger.error(f"Error reading .txt file {file_path}: {e}")
            return ""


class DocxExtractor(TextExtractor):
    """Text extractor for .docx files."""
    
    def can_extract(self, file_path: Path) -> bool:
        return file_path.suffix.lower() == '.docx'
    
    def extract_text(self, file_path: Path) -> str:
        """Extract text from .docx file using python-docx."""
        try:
            import docx
            doc = docx.Document(file_path)
            text = []
            for paragraph in doc.paragraphs:
                text.append(paragraph.text)
            return '\n'.join(text)
        except ImportError:
            logger.warning("python-docx not installed. Install with: pip install python-docx")
            return ""
        except Exception as e:
            logger.error(f"Error reading .docx file {file_path}: {e}")
            return ""


class DocExtractor(TextExtractor):
    """Text extractor for .doc files."""
    
    def can_extract(self, file_path: Path) -> bool:
        return file_path.suffix.lower() == '.doc'
    
    def extract_text(self, file_path: Path) -> str:
        """Extract text from .doc file using textract or python-docx fallback."""
        # Try textract first for .doc files
        try:
            import textract
            text = textract.process(str(file_path)).decode('utf-8')
            return text
        except ImportError:
            logger.warning("textract not installed. Trying alternative method...")
        except Exception as e:
            logger.error(f"Error reading .doc file {file_path} with textract: {e}")
        
        # Try python-docx as fallback
        try:
            import docx
            doc = docx.Document(file_path)
            text = []
            for paragraph in doc.paragraphs:
                text.append(paragraph.text)
            return '\n'.join(text)
        except ImportError:
            logger.error("Neither textract nor python-docx available for .doc files")
            return ""
        except Exception as e:
            logger.error(f"Fallback failed for .doc file {file_path}: {e}")
            return ""


class PdfExtractor(TextExtractor):
    """Text extractor for .pdf files."""
    
    def can_extract(self, file_path: Path) -> bool:
        return file_path.suffix.lower() == '.pdf'
    
    def extract_text(self, file_path: Path) -> str:
        """Extract text from .pdf file using PyPDF2 or textract fallback."""
        # Try PyPDF2 first
        try:
            import PyPDF2
            with open(file_path, 'rb') as file:
                reader = PyPDF2.PdfReader(file)
                text = []
                for page in reader.pages:
                    text.append(page.extract_text())
                return '\n'.join(text)
        except ImportError:
            logger.warning("PyPDF2 not installed. Trying textract for PDF...")
        except Exception as e:
            logger.error(f"Error reading PDF {file_path} with PyPDF2: {e}")
        
        # Fallback to textract
        try:
            import textract
            text = textract.process(str(file_path)).decode('utf-8')
            return text
        except ImportError:
            logger.error("Neither PyPDF2 nor textract available for PDF files")
            return ""
        except Exception as e:
            logger.error(f"Error reading PDF {file_path} with textract: {e}")
            return ""


class TextExtractionManager:
    """Manager class for text extraction from various document formats."""
    
    def __init__(self):
        self.extractors = [
            TxtExtractor(),
            DocxExtractor(),
            DocExtractor(),
            PdfExtractor()
        ]
    
    def extract_text(self, file_path: Path) -> str:
        """Extract text from file using appropriate extractor."""
        for extractor in self.extractors:
            if extractor.can_extract(file_path):
                return extractor.extract_text(file_path)
        
        logger.warning(f"No extractor available for file: {file_path}")
        return ""
    
    def get_supported_extensions(self) -> set:
        """Get set of supported file extensions."""
        return {'.txt', '.docx', '.doc', '.pdf'}
